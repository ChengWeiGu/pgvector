# -*- coding: utf-8 -*-
import os
import glob
import docx
import time
import argparse
from tqdm import tqdm
import DatabaseProcess
import EmbeddingFunction
from langchain.text_splitter import RecursiveCharacterTextSplitter



az_embed = EmbeddingFunction.AzureOpenAIEmbeddings()
pg_vector = DatabaseProcess.PGVector(DatabaseProcess.pg_setting)


'''**
*本範例暫不處理副檔名為.doc
*只處理.docx
*'''
class SpecScanner:
    
    @staticmethod
    def get_file_lists(directory):
        doc_files = [] #紀錄副檔名.doc
        docx_files = [] #紀錄副檔名.docx

        for file in glob.glob(os.path.join(directory, '**'), recursive=True):
            if file.endswith('.doc'):
                doc_files.append(file)
            elif file.endswith('.docx'):
                docx_files.append(file)  
        return doc_files, docx_files
    
    @staticmethod
    def get_filename_without_extension(path):
        filename_with_extension = os.path.basename(path)
        filename_without_extension, _ = os.path.splitext(filename_with_extension)
        return filename_without_extension
    
    
    def __init__(self, directory = "./SVN_datasheet", table_name = "spec"):
        # SVN dir name
        self.directory = directory
        # pg table name
        self.table_name = table_name
    
    '''**
    * 讀取文件並將表格資訊抽出，組成特定格式的chunk text
    * @params: file_path - 檔案路徑
    * @return: text_and_tables - 檔案特定格式的chunk text
    *'''
    def read_docx_file(self, file_path:str) -> str:
        doc = docx.Document(file_path)
        text_and_tables = ""
        text_and_tables += 'This is the '+ str(self.get_filename_without_extension(file_path)) +' Specification datasheet.'+ "\n"
        
        # Read paragraphs
        for para in doc.paragraphs:
            text_and_tables += para.text + "\n"

        # Read tables
        for table in doc.tables:
            text_and_tables += "Table: \n"
            for row in table.rows:
                row_data = "\t".join(cell.text for cell in row.cells)
                text_and_tables += row_data + "\n"

        return text_and_tables
    
    '''**
    * 建立 doc and docx 附檔名的列表
    * @params: NA
    * @return: NA
    * 有些文檔token數~15000 > 8192, 則split
    *'''
    def scan_folder_and_create_embed2pg(self):
        # 建立 text splitter 
        text_splitter_recur = RecursiveCharacterTextSplitter(
            chunk_size = 6000,
            chunk_overlap = 1200,
            length_function = len,
            is_separator_regex=False
        )
        # 獲取兩種不同附檔名的Word (此範例不處理doc_files)
        doc_files, docx_files = self.get_file_lists(self.directory)
        total_chunk_cnt = 0
        fail_list = []
        for idx, docx_file in tqdm(enumerate(docx_files), 
                                   desc="Process spec ",
                                   total=len(docx_files),
                                   unit="files"):
            try:
                text_and_tables = self.read_docx_file(docx_file)
                # 計算token數
                num_tokens = EmbeddingFunction.num_tokens_from_string_embed(text_and_tables)
                # 建立meta
                file_basename = os.path.basename(docx_file)
                split_key_word = '_Datasheet'
                if split_key_word not in file_basename:
                    split_key_word = '-Datasheet'
                hmi_model_name = file_basename.split(split_key_word)[0]
                metadata = {
                    'source': file_basename,
                    "model": hmi_model_name
                }
                # 處理metadata格式以便插入pg, 外掛 chunk_context and embedding
                data_keys = list(metadata.keys()) + ["chunk_context","embedding"]
                col_names = ", ".join(data_keys)
                # 當token數過大則split
                chunk_list = text_splitter_recur.split_text(text_and_tables) if num_tokens > 8192 else [text_and_tables]
                for chunk in chunk_list:
                    # 加入 chunk_context
                    data_values = list(metadata.values())
                    data_values.append(chunk)
                    # 轉換Embedding, 失敗試三次
                    for _ in range(3):
                        chunk_embedding = az_embed.get_embedding(chunk)
                        if chunk_embedding is not None:
                            data_values.append(chunk_embedding)
                            # 存入pg
                            data_values = [tuple(data_values)]
                            pg_vector.upsert_data(self.table_name, 
                                                col_names, 
                                                data_values)
                            total_chunk_cnt += 1
                            break
                        else:
                            time.sleep(10)
                
            except Exception as e:
                fail_list.append((docx_file, str(e)))
            
            
        print("total len of chunks: ",total_chunk_cnt)
        print("fail list:\n",fail_list)
    

def main():
    
    parser = argparse.ArgumentParser(description='Spec document')
    parser.add_argument('-t', '--table_name', type=str, default="spec", help='創建Postgres Table名稱存入spec')
    parser.add_argument('-s', '--src', type=str, default="./SVN_datasheet", help='Data Sheet 資料夾位置')
    args = parser.parse_args()
    
    table_name = args.table_name
    if table_name.strip() == "":
        print("table_name is empty")
        return
    
    directory = args.src
    if directory.strip() == "":
        print("directory is empty")
        return
    
    print(f"Create table '{table_name}' if not exists....")
    ret_json = pg_vector.create_spec_table(table_name=table_name)
    if ret_json["status"] == "fail":
        print(ret_json["error_reason"])
        return
    
    
    print("start creating datasheets....")
    scanner = SpecScanner(directory=directory, table_name=table_name)
    scanner.scan_folder_and_create_embed2pg()
    print("done")
    
        
        
if __name__ == "__main__":
    main()